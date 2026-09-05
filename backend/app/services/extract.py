from __future__ import annotations

from .. import i18n
import csv
import io
import json
import logging
from pathlib import Path

log = logging.getLogger("projectai.extract")

TEXT_EXT = {".txt", ".md", ".mdx", ".rst", ".log", ".json", ".yaml", ".yml", ".xml", ".html", ".htm"}
AUDIO_VIDEO_EXT = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".aac", ".mp4", ".mov", ".avi", ".mkv", ".webm"}


class ExtractError(RuntimeError):
    pass


def is_audio_video(filename: str) -> bool:
    return Path(filename).suffix.lower() in AUDIO_VIDEO_EXT


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        from charset_normalizer import from_bytes

        best = from_bytes(raw).best()
        if best is None:
            raise ExtractError(i18n._("Не удалось определить кодировку"))
        return str(best)


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:  # повреждённые страницы не валят весь файл
            log.warning("PDF %s: страница %d не извлеклась: %s", path.name, i, e)
            pages.append("")
    text = "\n\n".join(pages).strip()
    if not text:
        raise ExtractError(i18n._("PDF не содержит извлекаемого текста (возможно, скан — нужен OCR)"))
    return text


def _extract_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    text = "\n".join(p for p in parts if p and p.strip()).strip()
    if not text:
        raise ExtractError(i18n._("Документ пуст"))
    return text


def _extract_xlsx(path: Path) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    out = io.StringIO()
    for ws in wb.worksheets:
        out.write(i18n._("## Лист: {title}").format(title=ws.title) + "\n")
        writer = csv.writer(out, delimiter="\t")
        for row in ws.iter_rows(values_only=True):
            if row and any(c is not None for c in row):
                writer.writerow(["" if c is None else str(c) for c in row])
        out.write("\n")
    wb.close()
    return out.getvalue().strip()


def _extract_csv(path: Path) -> str:
    return _read_text(path)


def extract_text(path: str | Path) -> str:
    """Извлечение текста из документа. Для аудио/видео используй transcribe."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(p)
    if ext in {".docx", ".dotx"}:
        return _extract_docx(p)
    if ext == ".doc":
        raise ExtractError(i18n._("Формат .doc (старый Word) не поддерживается — сохрани как .docx"))
    if ext in {".xlsx", ".xlsm", ".xltx"}:
        return _extract_xlsx(p)
    if ext in {".csv", ".tsv"}:
        return _extract_csv(p)
    if ext in TEXT_EXT:
        text = _read_text(p)
        if ext == ".json":
            try:  # компактный json переформатируем для читаемости
                text = json.dumps(json.loads(text), ensure_ascii=False, indent=1)[:500_000]
            except json.JSONDecodeError:
                pass
        return text
    if is_audio_video(p.name):
        raise ExtractError(i18n._("Это аудио/видео — используется транскрибация, не извлечение текста"))
    raise ExtractError(i18n._("Неподдерживаемый формат: {ext}").format(ext=ext))
