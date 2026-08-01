from __future__ import annotations

from pathlib import Path

import pytest

from app.services.extract import ExtractError, extract_text, is_audio_video


def test_txt(tmp_path: Path) -> None:
    p = tmp_path / "note.txt"
    p.write_text("Привет, это заметка", encoding="utf-8")
    assert extract_text(p) == "Привет, это заметка"


def test_md_and_json(tmp_path: Path) -> None:
    md = tmp_path / "doc.md"
    md.write_text("# Заголовок\nтекст", encoding="utf-8")
    assert "Заголовок" in extract_text(md)
    js = tmp_path / "data.json"
    js.write_text('{"a":1}', encoding="utf-8")
    assert '"a"' in extract_text(js)


def test_cp1251(tmp_path: Path) -> None:
    p = tmp_path / "win.txt"
    p.write_bytes("Русский текст в cp1251".encode("cp1251"))
    assert "Русский текст" in extract_text(p)


def test_docx(tmp_path: Path) -> None:
    import docx

    doc = docx.Document()
    doc.add_paragraph("Первый абзац ТЗ")
    doc.add_paragraph("Второй абзац")
    p = tmp_path / "tz.docx"
    doc.save(str(p))
    text = extract_text(p)
    assert "Первый абзац ТЗ" in text and "Второй абзац" in text


def test_xlsx(tmp_path: Path) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Смета"
    ws.append(["Фича", "Оценка"])
    ws.append(["Авторизация", 8])
    p = tmp_path / "estimate.xlsx"
    wb.save(str(p))
    text = extract_text(p)
    assert "Смета" in text and "Авторизация" in text


def test_pdf_no_text(tmp_path: Path) -> None:
    from pypdf import PdfWriter

    w = PdfWriter()
    w.add_blank_page(width=200, height=200)
    p = tmp_path / "blank.pdf"
    with open(p, "wb") as f:
        w.write(f)
    with pytest.raises(ExtractError):
        extract_text(p)


def test_audio_video_detection() -> None:
    assert is_audio_video("call.m4a")
    assert is_audio_video("Recording.MP4")
    assert not is_audio_video("doc.pdf")
    with pytest.raises(ExtractError):
        extract_text("call.m4a")
